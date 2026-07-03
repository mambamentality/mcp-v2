import base64
import io
import json
import logging
from datetime import datetime, timezone
from pathlib import Path

import azure.functions as func
from docx import Document
from docx.shared import Pt

from .draft_manager import DraftManager
from .mcp_helpers import ToolProperty, get_arguments, tool_properties_json
from .schemas import Draft

logger = logging.getLogger(__name__)

_TOOL_PROPERTIES = tool_properties_json(
    [
        ToolProperty("draftId", "string", "Id del borrador aprobado a convertir en DOCX.", isRequired=True),
    ]
)


def register_generate_docx_tool(app: func.FunctionApp):
    manager = DraftManager()

    @app.mcp_tool_trigger(
        arg_name="context",
        tool_name="GenerateDocx",
        description="Convierte el borrador aprobado en un documento Word con estilos corporativos.",
        tool_properties=_TOOL_PROPERTIES,
    )
    def generate_docx(context) -> str:
        try:
            args = get_arguments(context)
            draft_id = args.get("draftId")
            if not draft_id:
                raise ValueError("Debe enviar 'draftId'.")

            draft = manager.get_draft(str(draft_id))
            if draft is None:
                raise ValueError(f"Borrador no encontrado: {draft_id}")
            if draft.status != "approved":
                raise ValueError("El borrador debe estar aprobado antes de generar el DOCX.")

            document = _build_document(draft)

            # Serializamos a memoria: el cliente MCP recibe el archivo en base64
            # dentro de la respuesta JSON, en vez de depender del filesystem local
            # (que en Azure Functions no es persistente ni compartido entre instancias).
            buffer = io.BytesIO()
            document.save(buffer)
            file_bytes = buffer.getvalue()
            file_name = f"Acta_{draft.draft_id}.docx"

            # Copia opcional a disco, solo útil para inspección en desarrollo local.
            _save_document_locally(file_bytes, file_name)

            return json.dumps(
                {
                    "status": "ok",
                    "fileName": file_name,
                    "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "fileBase64": base64.b64encode(file_bytes).decode("ascii"),
                },
                ensure_ascii=False,
            )
        except Exception as exc:
            logger.exception("GenerateDocx error")
            return json.dumps({"status": "error", "message": str(exc)}, ensure_ascii=False)

    return generate_docx


def _build_document(draft: Draft) -> Document:
    document = Document()
    styles = document.styles
    heading_style = styles["Heading 1"]
    heading_style.font.name = "Arial"
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True

    normal_style = styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    document.add_heading("Acta de Reunión", level=1)
    document.add_paragraph(f"Generado: {datetime.now(timezone.utc).isoformat()} UTC")
    document.add_paragraph("")

    for line in draft.draft_markdown.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("| ") and "|" in line[2:]:
            # Las tablas se procesan aparte en _insert_tables_from_markdown.
            continue
        else:
            document.add_paragraph(line)

    _insert_tables_from_markdown(document, draft.draft_markdown)

    if draft.approved_at:
        document.add_page_break()
        document.add_heading("Aprobación", level=2)
        document.add_paragraph(f"Aprobado el: {draft.approved_at.isoformat()}")

    return document


def _save_document_locally(file_bytes: bytes, file_name: str) -> None:
    try:
        output_dir = Path(__file__).resolve().parent.parent / ".data" / "docx"
        output_dir.mkdir(parents=True, exist_ok=True)
        (output_dir / file_name).write_bytes(file_bytes)
    except OSError:
        # En Azure el filesystem puede ser de solo lectura; no debe romper la tool.
        logger.warning("No se pudo escribir copia local del DOCX (filesystem no escribible).")


def _insert_tables_from_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    table_rows = []
    in_table = False

    for line in lines:
        if line.startswith("| "):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if not in_table:
                in_table = True
                table_rows = [cells]
            else:
                table_rows.append(cells)
        else:
            if in_table:
                _add_table(document, table_rows)
                in_table = False
                table_rows = []

    if in_table and table_rows:
        _add_table(document, table_rows)


def _add_table(document: Document, rows: list) -> None:
    if not rows:
        return
    table = document.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
